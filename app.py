import meraki
from prettytable import PrettyTable
from flask import session
from flask import Flask, render_template, flash, url_for, session, redirect, send_file, after_this_request
from flask_wtf import FlaskForm
from wtforms import RadioField, StringField, SubmitField
from wtforms.validators import DataRequired
from docx import Document
from docx.shared import Pt
import os


from GET_Templates import getTemplate
from GET_firmware import firmware
from GET_License import getLicence
from Get_WarmSpare import getWarmSpare
from GET_AlertsNetwork import getAlertsNetwork
from GET_NetworkDevices import getNetworkDevicesWeb
from Get_SwitchStacksNetwork import getSwitchStack
from GET_MS_mtu import mtu
from GET_MsDetails import getMsDetail
from GET_VpnBgp import getVpnBgp
from GET_Site2Site_VPN import getNetworkVpn
from GET_MxSetting import  getMxSettings
from GET_MxVlans import getNetMxVlan
from GET_snmp import getSnmp
from GET_Ssid import getSsid
from GET_rfProfiles import get_rfProfiles
from GET_l3Fw import getl3Fw
from GET_channelUtil import getChannelUtil

app = Flask(__name__)

app.config['SECRET_KEY'] = 'cisco'



@app.route('/', methods=['GET', 'POST']) #127.0.0.1:5000
def index():
    #Check if a file exist and delete it
    fp = 'Meraki Health Check.docx'
    if os.path.isfile(fp):
        os.remove(fp)
        print(f"{fp} deleted")
    else:
        print("File not found")

    class apiForm(FlaskForm):
        api = StringField("Please enter your Meraki API key: ", validators=[DataRequired()], render_kw={"class": "form-control"})
        submit = SubmitField('Submit', render_kw={"class": "btn btn-success w-100"})
    form = apiForm()
    if form.validate_on_submit():
        session['api'] = form.api.data
        return redirect(url_for('orgs'))
    else:
        if 'api' in session:
            return redirect(url_for('orgs'))

        return render_template('index.html', form=form)

@app.route('/orgs', methods = ['GET', 'POST'])
def orgs():
    if 'api' in session:
        api = session['api']
        dashboard = meraki.DashboardAPI(api, print_console=False, output_log=False)
        response = dashboard.organizations.getOrganizations()
        list_data = []
        for data in response:
            orgName = (data['name'])
            orgId = (data['id'])
            list_Temp = (orgId, orgName)
            list_data.append(list_Temp)
        list_data.sort()



        class OrgForm(FlaskForm):
            org = RadioField('Please choose a desire organization', choices=list_data, validators=[DataRequired()])
            submit = SubmitField('Next', render_kw={"class": "btn btn-success w-100"})
        form = OrgForm()
        print(form.validate_on_submit())
        if form.validate_on_submit():
            print(form.validate_on_submit())
            session['orgs'] = form.org.data
            return redirect(url_for('networks'))

        return render_template('orgs.html', list_data=list_data, form=form)
    else:
        return redirect(url_for('logout'))

@app.route('/networks', methods=['GET', 'POST'])
def networks():
    if 'api' and 'orgs' in session:
        api = session['api']
        orgId = session['orgs']
        dashboard = meraki.DashboardAPI(api, print_console=False, output_log=False)
        response = dashboard.organizations.getOrganizationNetworks(orgId, total_pages='all')
        list_data = []
        list_data.sort()
        for data in response:
            networkName = (data['name'])
            netId = (data['id'])
            list_Temp = (netId, networkName)
            list_data.append(list_Temp)
        list_data.sort()

        class NetworksForm(FlaskForm):
            networks = RadioField('Please choose a desire network', choices=list_data, validators=[DataRequired()])
            submit = SubmitField('Next', render_kw={"class": "btn btn-success w-100"})
        form = NetworksForm()
        if form.validate_on_submit():
            session['networks'] = form.networks.data
            selected_option = form.networks.data
            print(selected_option)
            return redirect(url_for('report'))

        return  render_template('networks.html', orgId=orgId, form=form)
    else:
        return redirect(url_for('logout'))

@app.route('/report', methods=['GET','POST'])
def report():
    if 'api' and 'orgs' and 'networks' in session:
        return render_template('report.html')
    else:
        return redirect(url_for('logout'))
    """
    api = session['api']
    network_id = session['networks']
    orgId = session['orgs']

    table_snmp = getSnmp(api, orgId)
    table_snmp = table_snmp.get_html_string(attributes={"class":"table table-sm table-striped table-bordered"})

    table_licenses = getLicence(api, orgId)
    table_licenses = table_licenses.get_html_string(attributes={"class":"table table-sm table-striped table-bordered"})

    table_templates = getTemplate(api, orgId)
    table_templates = table_templates.get_html_string(attributes={"class":"table table-sm table-striped table-bordered"})

    table_firmware = firmware(api, network_id)
    table_firmware = table_firmware.get_html_string(attributes={"class":"table table-sm table-striped table-bordered"})

    table_channelUtil = getChannelUtil(api, network_id)
    table_channelUtil = table_channelUtil.get_html_string(attributes={"class":"table table-sm table-striped table-bordered"})

    Add this in render_template if you want to add the output
    (table_snmp=table_snmp, table_licenses=table_licenses, table_templates=table_templates, table_channelUtil=table_channelUtil, table_firmware=table_firmware)
    """

@app.route('/download', methods=['GET','POST'])
def download_report():
    if 'api' and 'orgs' and 'networks' in session:
        api = session['api']
        network_id = session['networks']
        orgId = session['orgs']

        print(network_id)

        # Document
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(8)

        # Add Title
        document.add_heading('Meraki Health Check', 0)
        document.add_paragraph(' ')

        def addSection(table):
            # Add Header before the table
            document.add_heading(table.title, 4)
            document.add_paragraph(description)

            # Add a table to the document
            rows = len(table._rows) + 1
            cols = len(table.field_names)
            table_word = document.add_table(rows=rows, cols=cols, style='Light List Accent 1')

            # Add header row to the table
            for i, field in enumerate(table.field_names):
                cell = table_word.cell(0, i)
                cell.text = field

            # Add data rows to the table
            for i, row in enumerate(table._rows):
                for j, value in enumerate(row):
                    cell = table_word.cell(i + 1, j)
                    cell.text = str(value)

            document.add_paragraph(' ')

        document.add_heading("Organization Level", 1)

        # Organization Licenses
        table = getLicence(api, orgId)
        description = ("Below you can find the Organization licenses")
        addSection(table)

        # Organization Templates
        table = getTemplate(api, orgId)
        description = ("Below you can find the configured Organizations Templates ")
        addSection(table)

        # SNMP
        table = getSnmp(api, orgId)
        description = ("Organization snmp settings")
        addSection(table)

        document.add_heading("Network Level", 1)
        document.add_heading("Devices", 2)


        # Network Devices
        list_serial, table = getNetworkDevicesWeb(api, network_id)
        description = "Below you can find the devices list  in a network "
        addSection(table)

        # Firmware
        table = firmware(api, network_id)
        description = ("Below you can find the current firmware version per device")
        addSection(table)

        document.add_heading("Alerts", 2)

        # Network Alerts
        table1, table = getAlertsNetwork(api, network_id)
        description = ("Default alerts configuration on the network")
        addSection(table1)
        description = ("Default individual alerts configuration on the network")
        addSection(table)

        document.add_heading("MS(switches) info", 2)
        # Network Switches info
        table1, table = getMsDetail(api, list_serial)
        description = ('MS switch port statuses')
        addSection(table1)
        description = ("MS port details")
        addSection(table)

        # Network switch stack info
        table = getSwitchStack(api, network_id)
        description = "Below you can find the swtich stack information on the network"
        addSection(table)

        # Network MTU configuration
        table = mtu(api, network_id)
        description = "Recommended to keep at default of 9578 unless intermediate devices don’t support jumbo frames. " \
                      "This is useful to optimize server-to-server and application performance. Avoid fragmentation when possible."
        addSection(table)

        document.add_heading("MX(appliance) info", 2)

        # Network MX settings
        table = getMxSettings(api, network_id)
        description = ("The Cisco Meraki MX security appliance has a number of deployment options to meet the needs "
                       "of your network and infrastructure. Whether as the main edge firewall for your network, or as a concentrator device in your data center, the MX security appliance can be easily integrated.")
        addSection(table)

        # Network MX Vlans
        table = getNetMxVlan(api, network_id)
        description = ("List the VLANs for an MX network")
        addSection(table)

        # Network L3 Firewall Rules
        table = getl3Fw(api, network_id)
        description = ("Layer 3 Firewall rules configured on the network")
        addSection(table)

        # Network MX in warmspare
        table = getWarmSpare(api, network_id)
        description = ("Below you can find the MX Warmspare details on the network")
        addSection(table)

        document.add_heading("BGP", 3)
        # Network MX Hub BGP Configuration
        table1, table = getVpnBgp(api, network_id)
        p = document.add_paragraph("** Deployment mode should be configured as ")
        p.add_run('passthrough').bold = True

        description = ("MX appliances use iBGP to exchange route information over Meraki AutoVPN. "
                       "MXs deployed as one-armed VPN concentrator use eBGP to exchange and learn routes within the datacenter."
                       " Learned routes are redistributed to AutoVPN spokes using iBGP.")
        addSection(table)
        addSection(table1)

        document.add_heading("Site-to-Site VPN", 3)
        # Network site-to-site VPN settings
        table3, table2, table = getNetworkVpn(api, network_id)
        description = "Meraki Auto VPN technology is a unique solution that allows site-to-site VPN tunnel creation"
        addSection(table)
        addSection(table2)
        addSection(table3)

        document.add_heading("MR(wireless) info", 2)

        # ssid's
        table = getSsid(api, network_id)
        description = "List of SSIDs in the network"
        addSection(table)

        # RF profiles
        table = get_rfProfiles(api, network_id)
        description = "List of wireless Radio Frequency profiles in the network"
        addSection(table)

        # Channel Utilization
        table = getChannelUtil(api, network_id)
        description = "Channel utilization below 60% in 2.4GHz and below 40% in 5GHz in a 1 day timespan"
        addSection(table)

        document.add_page_break()

        # Save the document
        document.save('Meraki Health Check.docx')


        p = 'Meraki Health Check.docx'
        return send_file(p, as_attachment=True)
    else:
        return redirect(url_for('logout'))


@app.route('/logout')
def logout():
    session.pop('api', None)
    session.pop('orgs', None)
    session.pop('networks', None)
    return redirect(url_for('index'))



if __name__ == '__main__':
    app.run(debug=True)

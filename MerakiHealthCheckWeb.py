from docx import Document
from docx.shared import Pt
from docxcompose.composer import Composer

from GET_Templates import getTemplate
from GET_firmware import firmware
from GET_License import getLicence
from Get_WarmSpare import getWarmSpare
from GET_AlertsNetwork import getAlertsNetwork
from GET_NetworkDevices import getNetworkDevicesWeb
from Get_SwitchStacksNetwork import getSwitchStack
from GET_MS_mtu import mtu
from GET_MsDetails import getMsDetail
from GET_VpnBgp import getVpnBgp
from GET_Site2Site_VPN import getNetworkVpn
from GET_MxSetting import  getMxSettings
from GET_MxVlans import getNetMxVlan
from GET_snmp import getSnmp
from GET_Ssid import getSsid
from GET_rfProfiles import get_rfProfiles
from GET_l3Fw import getl3Fw
from GET_channelUtil import getChannelUtil
from GET_Network_Events import getNetworkEvents
from GET_Top_Ssid_by_Usage import get_top_ssid_by_usage
from GET_TopDevices_by_Usage import get_top_devices_by_usage
from Get_TopDevices_Model_by_Usage import get_top_devices_Model_by_usage
from GET_Top_Sw_by_Energy_Usage import get_top_sw_by_energy_usage

# Put the output in a Word document using python-docx"""
def create_document(api, orgId, network_id):
    # Document
    document = Document("CiscoTemplate.docx")
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Add Title
    #document.add_heading('Meraki Health Check', 0)
    document.add_paragraph(' ')

    def addSection(table):
        # Add Header before the table
        document.add_heading(table.title, 4)
        document.add_paragraph(description)

        # Add a table to the document
        rows = len(table._rows) + 1
        cols = len(table.field_names)
        table_word = document.add_table(rows=rows, cols=cols, style='Cisco CX Table | Default')

        # Add header row to the table
        for i, field in enumerate(table.field_names):
            cell = table_word.cell(0, i)
            cell.text = field

        # Add data rows to the table
        for i, row in enumerate(table._rows):
            for j, value in enumerate(row):
                cell = table_word.cell(i + 1, j)
                cell.text = str(value)

        document.add_paragraph(' ')

    document.add_heading("Organization Level", 1)

    # Organization Licenses
    table = getLicence(api, orgId)
    description = ("Below you can find the Organization licenses")
    addSection(table)

    # Organization Templates
    table = getTemplate(api, orgId)
    description = ("Below you can find the configured Organizations Templates ")
    addSection(table)

    # SNMP
    table = getSnmp(api, orgId)
    description = ("Organization snmp settings")
    addSection(table)

    #Top 10 Devices by Usage
    table = get_top_devices_by_usage(api, orgId)
    description = ("Top 10 devices by usage in the Organization (one month data)")
    addSection(table)

    # Top 10 Devices Model by Usage
    table = get_top_devices_Model_by_usage(api, orgId)
    description = ("Top 10 devices Model by usage in the Organization (one month data)")
    addSection(table)

    # Top SSIDs by Usage
    table = get_top_ssid_by_usage(api, orgId)
    description = "Top SSIDs by Usage in the Organization (one month data)"
    addSection(table)

    # Top 10 SW by Energy Usage
    table = get_top_sw_by_energy_usage(api, orgId)
    description = ("Top 10 Switches by Energy usage in the Organization (one month data)")
    addSection(table)

    document.add_heading("Network Level", 1)
    document.add_heading("Devices", 2)

    # Network Devices
    list_serial, table = getNetworkDevicesWeb(api, network_id)
    description = "Below you can find the devices list  in a network "
    addSection(table)

    # Network Events
    table = getNetworkEvents(api,network_id)
    description = "Below you can fin the top 10 events find in your network"
    addSection(table)

    # Firmware
    table = firmware(api, network_id)
    description = ("Below you can find the current firmware version per device")
    addSection(table)

    document.add_heading("Alerts", 2)

    # Network Alerts
    table1, table = getAlertsNetwork(api, network_id)
    description = ("Default alerts configuration on the network")
    addSection(table1)
    description = ("Default individual alerts configuration on the network")
    addSection(table)

    document.add_heading("MS(switches) info", 2)
    # Network Switches info
    table1, table = getMsDetail(api, list_serial)
    description = ('MS switch port statuses')
    addSection(table1)
    description = ("MS port details")
    addSection(table)

    # Network switch stack info
    table = getSwitchStack(api, network_id)
    description = "Below you can find the swtich stack information on the network"
    addSection(table)

    # Network MTU configuration
    table = mtu(api, network_id)
    description = "Recommended to keep at default of 9578 unless intermediate devices don’t support jumbo frames. " \
                  "This is useful to optimize server-to-server and application performance. Avoid fragmentation when possible."
    addSection(table)

    document.add_heading("MX(appliance) info", 2)

    # Network MX settings
    table = getMxSettings(api, network_id)
    description = ("The Cisco Meraki MX security appliance has a number of deployment options to meet the needs "
                   "of your network and infrastructure. Whether as the main edge firewall for your network, or as a concentrator device in your data center, the MX security appliance can be easily integrated.")
    addSection(table)

    # Network MX Vlans
    table = getNetMxVlan(api, network_id)
    description = ("List the VLANs for an MX network")
    addSection(table)

    # Network L3 Firewall Rules
    table = getl3Fw(api, network_id)
    description = ("Layer 3 Firewall rules configured on the network")
    addSection(table)

    # Network MX in warmspare
    table = getWarmSpare(api, network_id)
    description = ("Below you can find the MX Warmspare details on the network")
    addSection(table)

    document.add_heading("BGP", 3)
    # Network MX Hub BGP Configuration
    table1, table = getVpnBgp(api, network_id)
    p = document.add_paragraph("** Deployment mode should be configured as ")
    p.add_run('passthrough').bold = True

    description = ("MX appliances use iBGP to exchange route information over Meraki AutoVPN. "
                   "MXs deployed as one-armed VPN concentrator use eBGP to exchange and learn routes within the datacenter."
                   " Learned routes are redistributed to AutoVPN spokes using iBGP.")
    addSection(table)
    addSection(table1)

    document.add_heading("Site-to-Site VPN", 3)
    # Network site-to-site VPN settings
    table3, table2, table = getNetworkVpn(api, network_id)
    description = "Meraki Auto VPN technology is a unique solution that allows site-to-site VPN tunnel creation"
    addSection(table)
    addSection(table2)
    addSection(table3)

    document.add_heading("MR(wireless) info", 2)

    # ssid's
    table = getSsid(api, network_id)
    description = "List of SSIDs in the network"
    addSection(table)



    # RF profiles
    table = get_rfProfiles(api, network_id)
    description = "List of wireless Radio Frequency profiles in the network"
    addSection(table)

    # Channel Utilization
    table = getChannelUtil(api, network_id)
    description = "Channel utilization below 60% in 2.4GHz and below 40% in 5GHz in a 1 day timespan"
    addSection(table)

    document.add_page_break()

    # Save the document
    document.save('Meraki Health Check_Pre.docx')

    document2 = Document("CiscoTemplateEnd.docx")
    # name of the file you want to merge the docx file into
    document = Document("Meraki Health Check_Pre.docx")
    document = Composer(document)
    # name of the second docx file
    document2 = Document("CiscoTemplateEnd.docx")
    # append the doc2 into the master using append function
    document.append(document2)
    # Save the combined docx with a name
    document.save("Meraki Health Check.docx")

    print("")
    print("Document successfully created!")
